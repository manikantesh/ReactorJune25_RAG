#!/usr/bin/env python3
"""
Sample Document Generator for Legal AI Assistant

Creates sample legal documents in various formats (PDF, DOCX, TXT) for testing.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz  # PyMuPDF

def create_sample_docx():
    """Create a sample DOCX legal document."""
    doc = Document()
    
    # Title
    title = doc.add_heading('IN THE UNITED STATES DISTRICT COURT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_heading('FOR THE NORTHERN DISTRICT OF CALIFORNIA', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Case information
    doc.add_paragraph('Case No. 3:23-cv-01234')
    doc.add_paragraph('')
    
    # Parties
    doc.add_paragraph('ROBERT ANDERSON, Plaintiff,')
    doc.add_paragraph('v.')
    doc.add_paragraph('TECHNOLOGY CORP., Defendant.')
    doc.add_paragraph('')
    
    # Document type
    doc.add_heading('COMPLAINT FOR DAMAGES', 1)
    doc.add_paragraph('')
    
    # Introduction
    doc.add_paragraph('Plaintiff Robert Anderson ("Plaintiff") alleges as follows:')
    doc.add_paragraph('')
    
    # Jurisdiction
    doc.add_heading('JURISDICTION AND VENUE', 2)
    doc.add_paragraph('1. This Court has jurisdiction over this action pursuant to 28 U.S.C. § 1332 because the parties are citizens of different states and the amount in controversy exceeds $75,000.')
    doc.add_paragraph('')
    
    # Parties
    doc.add_heading('PARTIES', 2)
    doc.add_paragraph('2. Plaintiff Robert Anderson is a resident of California.')
    doc.add_paragraph('3. Defendant Technology Corp. is a Delaware corporation with its principal place of business in Texas.')
    doc.add_paragraph('')
    
    # Facts
    doc.add_heading('FACTS', 2)
    doc.add_paragraph('4. On January 15, 2023, Plaintiff entered into a software licensing agreement with Defendant.')
    doc.add_paragraph('5. Defendant failed to deliver the software as specified in the agreement.')
    doc.add_paragraph('6. Plaintiff suffered damages in excess of $100,000 as a result of Defendant\'s breach.')
    doc.add_paragraph('')
    
    # Causes of action
    doc.add_heading('CAUSES OF ACTION', 2)
    doc.add_paragraph('FIRST CAUSE OF ACTION - BREACH OF CONTRACT')
    doc.add_paragraph('7. Plaintiff incorporates by reference paragraphs 1-6 above.')
    doc.add_paragraph('8. Defendant breached the software licensing agreement by failing to deliver the software.')
    doc.add_paragraph('9. Plaintiff is entitled to damages in the amount of $100,000.')
    doc.add_paragraph('')
    
    # Prayer for relief
    doc.add_heading('PRAYER FOR RELIEF', 2)
    doc.add_paragraph('WHEREFORE, Plaintiff prays for judgment as follows:')
    doc.add_paragraph('1. For damages in the amount of $100,000;')
    doc.add_paragraph('2. For costs of suit; and')
    doc.add_paragraph('3. For such other and further relief as the Court deems just and proper.')
    doc.add_paragraph('')
    
    # Signature
    doc.add_paragraph('DATED: March 15, 2023')
    doc.add_paragraph('')
    doc.add_paragraph('Respectfully submitted,')
    doc.add_paragraph('')
    doc.add_paragraph('JANE SMITH')
    doc.add_paragraph('Attorney for Plaintiff')
    doc.add_paragraph('123 Legal Street')
    doc.add_paragraph('San Francisco, CA 94102')
    doc.add_paragraph('Tel: (415) 555-0123')
    
    # Save the document
    output_path = Path('data/sample_documents/sample_complaint.docx')
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"✅ Created: {output_path}")

def create_sample_pdf():
    """Create a sample PDF legal document."""
    doc = fitz.open()
    page = doc.new_page()
    
    # Document content
    content = """
IN THE COURT OF APPEAL OF THE STATE OF CALIFORNIA

Case No. A123456

OPINION

JAMES WILSON, Appellant,
v.
CITY OF SAN DIEGO, Respondent.

Filed: February 20, 2023

OPINION BY JUSTICE GARCIA

This appeal presents the question of whether the trial court erred in granting summary judgment in favor of the City of San Diego in a personal injury action arising from a slip and fall on a public sidewalk.

FACTS

On March 10, 2022, appellant James Wilson was walking on a public sidewalk in downtown San Diego when he slipped and fell on a patch of ice. The temperature was below freezing, and the City had not applied salt or other de-icing materials to the sidewalk.

Wilson suffered a broken hip and required surgery. His medical expenses totaled $50,000, and he lost $20,000 in wages due to his inability to work.

PROCEDURAL HISTORY

Wilson filed a complaint against the City alleging negligence in failing to maintain the sidewalk in a safe condition. The City moved for summary judgment, arguing that it had no duty to remove natural accumulations of ice and snow from public sidewalks.

The trial court granted the City's motion, finding that the City owed no duty to Wilson under the circumstances.

LEGAL ANALYSIS

The issue before this Court is whether the City had a duty to maintain the sidewalk in a safe condition during winter weather conditions.

Under California law, municipalities have a duty to maintain public property in a reasonably safe condition. However, this duty does not extend to removing natural accumulations of ice and snow unless the municipality has created an artificial condition that increases the risk of injury.

In this case, the ice on the sidewalk was a natural accumulation resulting from freezing temperatures. The City did not create any artificial condition that increased the risk of injury. Therefore, the City had no duty to remove the ice.

CONCLUSION

The trial court correctly granted summary judgment in favor of the City. The judgment is affirmed.

JUSTICE GARCIA
February 20, 2023
"""
    
    # Insert text into the page
    page.insert_text((50, 50), content, fontsize=12)
    
    # Save the PDF
    output_path = Path('data/sample_documents/sample_appeal_opinion.pdf')
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    doc.close()
    print(f"✅ Created: {output_path}")

def create_sample_contract_txt():
    """Create a sample contract in text format."""
    content = """SOFTWARE DEVELOPMENT AGREEMENT

This Software Development Agreement (the "Agreement") is entered into as of March 1, 2023, by and between:

DEVELOPMENT CORP., a California corporation ("Developer")
and
CLIENT INC., a Delaware corporation ("Client")

1. SERVICES

Developer agrees to develop custom software for Client according to the specifications set forth in Exhibit A.

2. PAYMENT

Client agrees to pay Developer $50,000 for the development services, payable as follows:
- $25,000 upon execution of this Agreement
- $25,000 upon delivery of the completed software

3. DELIVERY

Developer shall deliver the completed software to Client by June 1, 2023.

4. WARRANTIES

Developer warrants that the software will function according to the specifications for a period of one year from delivery.

5. TERMINATION

Either party may terminate this Agreement upon 30 days written notice.

6. GOVERNING LAW

This Agreement shall be governed by the laws of the State of California.

IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.

DEVELOPMENT CORP.

By: _________________
Name: John Developer
Title: CEO

CLIENT INC.

By: _________________
Name: Mary Client
Title: CTO

Date: March 1, 2023
"""
    
    output_path = Path('data/sample_documents/sample_contract.txt')
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    with open(output_path, 'w') as f:
        f.write(content)
    
    print(f"✅ Created: {output_path}")

def create_sample_employment_agreement_docx():
    """Create a sample employment agreement in DOCX format."""
    doc = Document()
    
    # Title
    title = doc.add_heading('EMPLOYMENT AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    doc.add_paragraph('This Employment Agreement (the "Agreement") is made and entered into as of April 1, 2023, by and between:')
    doc.add_paragraph('')
    doc.add_paragraph('STARTUP TECH INC., a California corporation ("Company")')
    doc.add_paragraph('and')
    doc.add_paragraph('SARAH JOHNSON ("Employee")')
    doc.add_paragraph('')
    
    # Terms
    doc.add_heading('1. POSITION AND DUTIES', 2)
    doc.add_paragraph('Company hereby employs Employee as Senior Software Engineer. Employee shall perform such duties as are customarily performed by a Senior Software Engineer and such other duties as may be assigned by Company.')
    doc.add_paragraph('')
    
    doc.add_heading('2. COMPENSATION', 2)
    doc.add_paragraph('Employee shall receive an annual salary of $120,000, payable in accordance with Company\'s normal payroll practices.')
    doc.add_paragraph('')
    
    doc.add_heading('3. BENEFITS', 2)
    doc.add_paragraph('Employee shall be eligible to participate in Company\'s employee benefit plans, including health insurance, 401(k), and paid time off.')
    doc.add_paragraph('')
    
    doc.add_heading('4. TERM', 2)
    doc.add_paragraph('This Agreement shall commence on April 1, 2023, and shall continue until terminated by either party in accordance with Section 5.')
    doc.add_paragraph('')
    
    doc.add_heading('5. TERMINATION', 2)
    doc.add_paragraph('Either party may terminate this Agreement at any time with 30 days written notice.')
    doc.add_paragraph('')
    
    doc.add_heading('6. CONFIDENTIALITY', 2)
    doc.add_paragraph('Employee agrees to maintain the confidentiality of Company\'s proprietary information during and after employment.')
    doc.add_paragraph('')
    
    doc.add_heading('7. GOVERNING LAW', 2)
    doc.add_paragraph('This Agreement shall be governed by the laws of the State of California.')
    doc.add_paragraph('')
    
    # Signature block
    doc.add_paragraph('IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.')
    doc.add_paragraph('')
    doc.add_paragraph('STARTUP TECH INC.')
    doc.add_paragraph('')
    doc.add_paragraph('By: _________________')
    doc.add_paragraph('Name: Michael CEO')
    doc.add_paragraph('Title: Chief Executive Officer')
    doc.add_paragraph('')
    doc.add_paragraph('SARAH JOHNSON')
    doc.add_paragraph('')
    doc.add_paragraph('Signature: _________________')
    doc.add_paragraph('Date: April 1, 2023')
    
    # Save the document
    output_path = Path('data/sample_documents/sample_employment_agreement.docx')
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"✅ Created: {output_path}")

def create_sample_criminal_complaint_pdf():
    """Create a sample criminal complaint in PDF format."""
    doc = fitz.open()
    page = doc.new_page()
    
    content = """
IN THE SUPERIOR COURT OF CALIFORNIA
COUNTY OF LOS ANGELES

Case No. CR-2023-567

THE PEOPLE OF THE STATE OF CALIFORNIA, Plaintiff,
v.
DAVID BROWN, Defendant.

CRIMINAL COMPLAINT

I, the undersigned, being duly sworn, state that I have probable cause to believe and do believe that the defendant committed the following offense(s):

COUNT 1: GRAND THEFT (Penal Code § 487(a))

On or about January 15, 2023, in the County of Los Angeles, State of California, the defendant DAVID BROWN did unlawfully take personal property, to wit: laptop computer, cellular phone, and cash, of a value exceeding $950, belonging to ABC Corporation.

COUNT 2: BURGLARY (Penal Code § 459)

On or about January 15, 2023, in the County of Los Angeles, State of California, the defendant DAVID BROWN did unlawfully enter a commercial building with the intent to commit theft.

FACTS SUPPORTING PROBABLE CAUSE:

On January 15, 2023, at approximately 2:30 AM, security cameras at ABC Corporation's office building captured the defendant entering the building through a rear door. The defendant was observed taking a laptop computer, cellular phone, and $500 in cash from an office.

The defendant was identified through security footage and was arrested on January 16, 2023, in possession of the stolen items.

WHEREFORE, I request that a warrant be issued for the arrest of the defendant.

DATED: January 17, 2023

Detective John Smith
Los Angeles Police Department
Badge No. 12345

Sworn to before me this 17th day of January, 2023.

Judge Maria Rodriguez
Superior Court of California
County of Los Angeles
"""
    
    # Insert text into the page
    page.insert_text((50, 50), content, fontsize=11)
    
    # Save the PDF
    output_path = Path('data/sample_documents/sample_criminal_complaint.pdf')
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    doc.close()
    print(f"✅ Created: {output_path}")

def main():
    """Create all sample documents."""
    print("🚀 Creating Sample Legal Documents in Multiple Formats...")
    print("=" * 60)
    
    # Create documents in different formats
    create_sample_docx()
    create_sample_pdf()
    create_sample_contract_txt()
    create_sample_employment_agreement_docx()
    create_sample_criminal_complaint_pdf()
    
    print("\n✅ All sample documents created successfully!")
    print("\n📁 Created files:")
    print("  - sample_complaint.docx (Civil complaint)")
    print("  - sample_appeal_opinion.pdf (Appellate court opinion)")
    print("  - sample_contract.txt (Software development contract)")
    print("  - sample_employment_agreement.docx (Employment agreement)")
    print("  - sample_criminal_complaint.pdf (Criminal complaint)")
    
    print("\n💡 You can now test document ingestion with these files:")
    print("  python3 src/utils/populate_database.py")

if __name__ == "__main__":
    main() 