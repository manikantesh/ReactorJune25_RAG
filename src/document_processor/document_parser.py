"""
Document Parser for Legal AI Assistant

Handles parsing of various legal document formats including PDF, DOCX, and text files.
Extracts structured information from legal documents for AI analysis.
"""

import os
import re
from typing import Dict, List, Optional, Any
from pathlib import Path
import logging

import PyPDF2
import pdfplumber
from docx import Document
import pytesseract
from PIL import Image
import fitz  # PyMuPDF

logger = logging.getLogger(__name__)


class DocumentParser:
    """Main document parser for legal documents."""
    
    def __init__(self, config: Optional[Dict] = None):
        """Initialize the document parser.
        
        Args:
            config: Configuration dictionary for parser settings
        """
        self.config = config or {}
        self.supported_formats = ['.pdf', '.docx', '.txt', '.jpg', '.jpeg', '.png']
        
    def parse_document(self, file_path: str) -> Dict[str, Any]:
        """Parse a legal document and extract structured information.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing parsed document information
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
            
        if file_path.suffix.lower() not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
            
        logger.info(f"Parsing document: {file_path}")
        
        try:
            if file_path.suffix.lower() == '.pdf':
                return self._parse_pdf(file_path)
            elif file_path.suffix.lower() == '.docx':
                return self._parse_docx(file_path)
            elif file_path.suffix.lower() == '.txt':
                return self._parse_text(file_path)
            elif file_path.suffix.lower() in ['.jpg', '.jpeg', '.png']:
                return self._parse_image(file_path)
            else:
                raise ValueError(f"Unsupported format: {file_path.suffix}")
                
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {str(e)}")
            raise
    
    def _parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Parse PDF documents using multiple methods for better extraction."""
        content = ""
        metadata = {}
        
        # Try PyMuPDF first for better text extraction
        try:
            doc = fitz.open(file_path)
            content = ""
            for page in doc:
                content += page.get_text()
            metadata = doc.metadata
            doc.close()
        except Exception as e:
            logger.warning(f"PyMuPDF failed, trying pdfplumber: {e}")
            
            # Fallback to pdfplumber
            try:
                with pdfplumber.open(file_path) as pdf:
                    content = ""
                    for page in pdf.pages:
                        if page.extract_text():
                            content += page.extract_text() + "\n"
                    metadata = pdf.metadata
            except Exception as e2:
                logger.warning(f"pdfplumber failed, trying PyPDF2: {e2}")
                
                # Final fallback to PyPDF2
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    content = ""
                    for page in reader.pages:
                        content += page.extract_text() + "\n"
                    metadata = reader.metadata
        
        return {
            'file_path': str(file_path),
            'file_type': 'pdf',
            'content': content,
            'metadata': metadata,
            'extracted_info': self._extract_legal_info(content)
        }
    
    def _parse_docx(self, file_path: Path) -> Dict[str, Any]:
        """Parse DOCX documents."""
        doc = Document(file_path)
        content = ""
        
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
            
        # Extract tables
        tables_content = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    tables_content += cell.text + " "
                tables_content += "\n"
        
        content += tables_content
        
        return {
            'file_path': str(file_path),
            'file_type': 'docx',
            'content': content,
            'metadata': {
                'title': doc.core_properties.title,
                'author': doc.core_properties.author,
                'created': str(doc.core_properties.created),
                'modified': str(doc.core_properties.modified)
            },
            'extracted_info': self._extract_legal_info(content)
        }
    
    def _parse_text(self, file_path: Path) -> Dict[str, Any]:
        """Parse plain text documents."""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            
        return {
            'file_path': str(file_path),
            'file_type': 'txt',
            'content': content,
            'metadata': {},
            'extracted_info': self._extract_legal_info(content)
        }
    
    def _parse_image(self, file_path: Path) -> Dict[str, Any]:
        """Parse image documents using OCR."""
        try:
            image = Image.open(file_path)
            content = pytesseract.image_to_string(image)
            
            return {
                'file_path': str(file_path),
                'file_type': 'image',
                'content': content,
                'metadata': {
                    'image_size': image.size,
                    'image_mode': image.mode
                },
                'extracted_info': self._extract_legal_info(content)
            }
        except Exception as e:
            logger.error(f"OCR failed for {file_path}: {e}")
            raise
    
    def _extract_legal_info(self, content: str) -> Dict[str, Any]:
        """Extract structured legal information from document content."""
        extracted = {
            'case_name': self._extract_case_name(content),
            'court': self._extract_court_info(content),
            'date': self._extract_date(content),
            'judges': self._extract_judges(content),
            'parties': self._extract_parties(content),
            'citation': self._extract_citation(content),
            'key_facts': self._extract_key_facts(content),
            'legal_issues': self._extract_legal_issues(content),
            'holding': self._extract_holding(content),
            'reasoning': self._extract_reasoning(content)
        }
        
        return extracted
    
    def _extract_case_name(self, content: str) -> Optional[str]:
        """Extract case name from document content."""
        # Common patterns for case names
        patterns = [
            r'(?:IN RE|In re|in re)\s+([A-Z][A-Z\s&,\.]+)',
            r'([A-Z][A-Z\s&,\.]+)\s+v\.\s+([A-Z][A-Z\s&,\.]+)',
            r'([A-Z][A-Z\s&,\.]+)\s+vs\.\s+([A-Z][A-Z\s&,\.]+)',
            r'Case\s+No\.\s*[:\-]?\s*([A-Z0-9\-]+)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, content[:2000])  # Search first 2000 chars
            if match:
                return match.group(0).strip()
        
        return None
    
    def _extract_court_info(self, content: str) -> Optional[str]:
        """Extract court information from document content."""
        court_patterns = [
            r'(Supreme Court|Court of Appeals|District Court|Circuit Court)',
            r'(Federal|State|County)\s+Court',
            r'([A-Z][a-z]+)\s+(Supreme|Appellate|District)\s+Court',
        ]
        
        for pattern in court_patterns:
            match = re.search(pattern, content[:2000])
            if match:
                return match.group(0).strip()
        
        return None
    
    def _extract_date(self, content: str) -> Optional[str]:
        """Extract date information from document content."""
        date_patterns = [
            r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
            r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b',
            r'\b\d{4}-\d{2}-\d{2}\b',
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, content[:2000])
            if match:
                return match.group(0).strip()
        
        return None
    
    def _extract_judges(self, content: str) -> List[str]:
        """Extract judge names from document content."""
        judge_patterns = [
            r'(?:Judge|Justice|Chief Justice)\s+([A-Z][a-z]+\s+[A-Z][a-z]+)',
            r'([A-Z][a-z]+\s+[A-Z][a-z]+),\s+(?:Judge|Justice)',
        ]
        
        judges = []
        for pattern in judge_patterns:
            matches = re.findall(pattern, content[:3000])
            judges.extend(matches)
        
        return list(set(judges))  # Remove duplicates
    
    def _extract_parties(self, content: str) -> List[str]:
        """Extract party names from document content."""
        party_patterns = [
            r'(?:Plaintiff|Defendant|Appellant|Appellee|Petitioner|Respondent)[:\s]+([A-Z][A-Z\s&,\.]+)',
        ]
        
        parties = []
        for pattern in party_patterns:
            matches = re.findall(pattern, content[:3000])
            parties.extend(matches)
        
        return list(set(parties))
    
    def _extract_citation(self, content: str) -> Optional[str]:
        """Extract legal citation from document content."""
        citation_patterns = [
            r'\d+\s+[A-Z]+\s+\d+',
            r'\d+\s+[A-Z]+\s+\d+\s+\(\d{4}\)',
            r'[A-Z]+\s+\d+\s+[A-Z]+\s+\d+',
        ]
        
        for pattern in citation_patterns:
            match = re.search(pattern, content[:2000])
            if match:
                return match.group(0).strip()
        
        return None
    
    def _extract_key_facts(self, content: str) -> List[str]:
        """Extract key facts from document content."""
        # This is a simplified extraction - in practice, this would be more sophisticated
        sentences = re.split(r'[.!?]+', content)
        fact_keywords = ['alleged', 'evidence', 'witness', 'testimony', 'found', 'determined']
        
        key_facts = []
        for sentence in sentences[:50]:  # Check first 50 sentences
            if any(keyword in sentence.lower() for keyword in fact_keywords):
                if len(sentence.strip()) > 20:  # Minimum length
                    key_facts.append(sentence.strip())
        
        return key_facts[:10]  # Return top 10 facts
    
    def _extract_legal_issues(self, content: str) -> List[str]:
        """Extract legal issues from document content."""
        issue_keywords = ['issue', 'question', 'whether', 'claim', 'cause of action']
        
        sentences = re.split(r'[.!?]+', content)
        legal_issues = []
        
        for sentence in sentences[:50]:
            if any(keyword in sentence.lower() for keyword in issue_keywords):
                if len(sentence.strip()) > 20:
                    legal_issues.append(sentence.strip())
        
        return legal_issues[:5]  # Return top 5 issues
    
    def _extract_holding(self, content: str) -> Optional[str]:
        """Extract the holding/decision from document content."""
        holding_keywords = ['hold', 'holding', 'conclude', 'find', 'determine', 'rule']
        
        sentences = re.split(r'[.!?]+', content)
        
        for sentence in sentences:
            if any(keyword in sentence.lower() for keyword in holding_keywords):
                if len(sentence.strip()) > 30:
                    return sentence.strip()
        
        return None
    
    def _extract_reasoning(self, content: str) -> List[str]:
        """Extract legal reasoning from document content."""
        reasoning_keywords = ['because', 'therefore', 'thus', 'consequently', 'reasoning']
        
        sentences = re.split(r'[.!?]+', content)
        reasoning = []
        
        for sentence in sentences:
            if any(keyword in sentence.lower() for keyword in reasoning_keywords):
                if len(sentence.strip()) > 30:
                    reasoning.append(sentence.strip())
        
        return reasoning[:5]  # Return top 5 reasoning statements 